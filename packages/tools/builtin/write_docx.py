"""write_docx — generate a Word document from structured sections."""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from pydantic import BaseModel, Field

from citnega.packages.protocol.callables.base import BaseCallable
from citnega.packages.protocol.callables.types import CallableType
from citnega.packages.tools.builtin._tool_base import ToolOutput, tool_policy

if TYPE_CHECKING:
    from citnega.packages.protocol.callables.context import CallContext


class DocxSection(BaseModel):
    heading: str = Field(default="", description="Section heading (empty for body-only paragraph).")
    body: str = Field(description="Section body text.")
    heading_level: int = Field(default=1, description="Heading level 1–3.")


class WriteDocxInput(BaseModel):
    title: str = Field(description="Document title (Heading 0 / Title style).")
    sections: list[DocxSection] = Field(description="Ordered sections to include.")
    filename: str = Field(description="Output file path, e.g. ~/Desktop/report.docx")
    author: str = Field(default="citnega", description="Core property author field.")


class WriteDocxTool(BaseCallable):
    """Generate a .docx Word document from a title and ordered sections."""

    name = "write_docx"
    description = (
        "Create a Microsoft Word (.docx) document from a title and list of sections. "
        "Each section has an optional heading (level 1–3) and body text. "
        "Returns the path of the created file."
    )
    callable_type = CallableType.TOOL
    input_schema = WriteDocxInput
    output_schema = ToolOutput
    policy = tool_policy(
        timeout_seconds=30.0,
        requires_approval=False,
        network_allowed=False,
    )

    async def _execute(self, input: WriteDocxInput, context: CallContext) -> ToolOutput:
        try:
            from docx import Document  # type: ignore[import-untyped]
            from docx.opc.constants import RELATIONSHIP_TYPE as RT  # noqa: F401
        except ImportError:
            return ToolOutput(result="[write_docx: python-docx not installed — run: pip install python-docx]")

        out_path = Path(input.filename).expanduser().resolve()
        out_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.core_properties.author = input.author
        doc.add_heading(input.title, level=0)

        for section in input.sections:
            if section.heading:
                level = max(1, min(3, section.heading_level))
                doc.add_heading(section.heading, level=level)
            if section.body:
                doc.add_paragraph(section.body)

        try:
            doc.save(str(out_path))
        except Exception as exc:
            return ToolOutput(result=f"[write_docx: failed to write: {exc}]")

        return ToolOutput(result=f"DOCX created: {out_path}  ({out_path.stat().st_size} bytes)")
